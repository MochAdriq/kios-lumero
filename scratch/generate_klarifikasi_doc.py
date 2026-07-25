from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Initialize Document
doc = Document()

# Set styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_paragraph('SURAT PENJELASAN DAN KLARIFIKASI ENTITAS USAHA')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.bold = True
    run.font.size = Pt(14)

doc.add_paragraph() # spacing

# Recipient
p = doc.add_paragraph('Kepada Yth,\nTim Verifikasi Midtrans\ndi Tempat')

doc.add_paragraph() # spacing

# Opening
doc.add_paragraph('Dengan hormat,')

p = doc.add_paragraph(
    'Sehubungan dengan catatan dari tim verifikasi Midtrans pada proses onboarding mengenai permintaan konfirmasi hubungan pengguna dengan entitas "PT Lokapedia Karya Bersama", bersama surat ini saya bermaksud menyampaikan klarifikasi dan koreksi data sebagai berikut:'
)

# Points
p1 = doc.add_paragraph(style='List Number')
p1.add_run('Bahwa terdapat koreksi penamaan entitas usaha. Nama entitas usaha yang benar dan sah yang saya gunakan dan daftarkan adalah ')
p1.add_run('PT. Lokapedia Sukses Bersama').bold = True
p1.add_run(', BUKAN PT Lokapedia Karya Bersama.')

p2 = doc.add_paragraph(style='List Number')
p2.add_run('Bahwa ')
p2.add_run('PT. Lokapedia Sukses Bersama').bold = True
p2.add_run(' merupakan entitas berbadan hukum yang berbentuk Perseroan Terbatas Perorangan (PT Perorangan). Entitas ini didirikan, dikelola, dan dimiliki sepenuhnya oleh saya selaku pengguna (merchant) yang mendaftarkan akun Midtrans ini.')

p3 = doc.add_paragraph(style='List Number')
p3.add_run('Bahwa seluruh aktivitas transaksi pembayaran, operasional platform Kios Lumero, serta pengelolaan website/sistem yang terintegrasi dengan Midtrans sepenuhnya merupakan tanggung jawab dari PT. Lokapedia Sukses Bersama.')

doc.add_paragraph() # spacing

# Closing
p = doc.add_paragraph('Demikian surat penjelasan dan klarifikasi ini saya buat dengan sebenar-benarnya dan penuh tanggung jawab, agar dapat dijadikan dasar pertimbangan untuk melanjutkan proses onboarding akun Midtrans kami. Atas perhatian dan kerjasamanya, saya ucapkan terima kasih.')

doc.add_paragraph() # spacing
doc.add_paragraph() # spacing

# Date and Signature
date_str = datetime.now().strftime('%d %B %Y')
p_sig = doc.add_paragraph(f'Hormat saya,\n\n\n\n\n\n________________________\nDirektur / Pendiri\nPT. Lokapedia Sukses Bersama')
p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Save document
output_path = 'Surat_Klarifikasi_Entitas_Midtrans.docx'
doc.save(output_path)
print(f"Dokumen berhasil dibuat: {output_path}")

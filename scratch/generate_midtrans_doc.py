"""
Generate DOCX: Alur Transaksi Pembelian Online & Panduan Uji Coba Pembayaran QRIS
Untuk Tim Verifikator Midtrans
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ──── CONFIG ────
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, '..', 'Panduan_Alur_Transaksi_QRIS_Kios_Lumero_v2.docx')
FLOW_DIAGRAM = r"C:\Users\HYPE R Series\.gemini\antigravity-ide\brain\cdaa987f-cb81-4d16-a44f-525ead3f5a1d\user_flow_diagram_1784652731271.png"

# Colors
NAVY = RGBColor(0x1a, 0x1a, 0x2e)
GREEN = RGBColor(0x34, 0xd3, 0x99)
DARK_GREEN = RGBColor(0x10, 0xb9, 0x81)
RED = RGBColor(0xff, 0x2d, 0x55)
GRAY = RGBColor(0x6b, 0x72, 0x80)
WHITE = RGBColor(0xff, 0xff, 0xff)
LIGHT_BG = RGBColor(0xf8, 0xfa, 0xfc)

doc = Document()

# ──── GLOBAL STYLE ────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35

# Heading styles
for level, (size, color) in enumerate([(24, NAVY), (16, NAVY), (13, NAVY)], 1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ──── HELPER FUNCTIONS ────
def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E2E8F0"/></w:pBdr>')
    pPr.append(pBdr)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" w:sz="{val.get("sz","4")}" w:space="0" w:color="{val.get("color","E2E8F0")}"/>')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1a1a2e')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            bg = 'F8FAFC' if r_idx % 2 == 0 else 'FFFFFF'
            set_cell_shading(cell, bg)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_screenshot_placeholder(doc, number, instruction):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create a bordered box effect
    run = p.add_run(f'📸 SCREENSHOT {number}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RED
    run.font.name = 'Calibri'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(16)
    run2 = p2.add_run(instruction)
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY
    run2.font.name = 'Calibri'

def add_step_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    
    run_num = p.add_run(f'Langkah {number}: ')
    run_num.bold = True
    run_num.font.size = Pt(12)
    run_num.font.color.rgb = GREEN if number == 7 else NAVY
    run_num.font.name = 'Calibri'
    
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = GREEN if number == 7 else NAVY
    run_title.font.name = 'Calibri'

# ═══════════════════════════════════════════
#  PAGE 1: COVER / TITLE PAGE
# ═══════════════════════════════════════════

# Spacer
for _ in range(4):
    doc.add_paragraph('')

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Alur Transaksi Pembelian Online\n& Panduan Uji Coba Pembayaran QRIS')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(8)
run = sub.add_run('Kios Lumero — Self-Order & Online Order Platform')
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

add_horizontal_line(doc)

# Meta info
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(16)
for line in [
    'Dokumen Verifikasi untuk Tim Midtrans',
    'Lingkungan: Sandbox Environment',
    'Versi Dokumen: 1.0',
    'Tanggal: Juli 2026',
]:
    run = meta.add_run(line + '\n')
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'

# Company info
company = doc.add_paragraph()
company.alignment = WD_ALIGN_PARAGRAPH.CENTER
company.paragraph_format.space_before = Pt(24)
run = company.add_run('PT. Lokapedia Sukses Bersama')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

run2 = company.add_run('\nhttps://lokapedia.id/lumero')
run2.font.size = Pt(10)
run2.font.color.rgb = GRAY
run2.font.name = 'Calibri'

# ═══════════════════════════════════════════
#  PAGE 2: DAFTAR ISI
# ═══════════════════════════════════════════
doc.add_page_break()

doc.add_heading('Daftar Isi', level=1)
add_horizontal_line(doc)

toc_items = [
    ('1.', 'Informasi Umum & Akses Platform'),
    ('2.', 'Kredensial Akun Uji Coba (Dummy)'),
    ('3.', 'Diagram Alur Transaksi (User Flow)'),
    ('4.', 'Detail Langkah-Langkah Pemesanan'),
    ('   4.1', 'Akses Halaman Welcome'),
    ('   4.2', 'Deteksi & Pemilihan Cabang (Outlet)'),
    ('   4.3', 'Pemilihan Menu & Keranjang (Cart)'),
    ('   4.4', 'Checkout & Konfigurasi Logistik'),
    ('   4.5', 'Pembayaran QRIS via Midtrans'),
    ('   4.6', 'Simulasi Pembayaran (Sandbox)'),
    ('   4.7', 'Notifikasi Sukses Real-time'),
    ('5.', 'Alur Verifikasi Backend & Webhook'),
    ('6.', 'Penutup'),
]

for num, label in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_n = p.add_run(f'{num}  ')
    run_n.bold = True
    run_n.font.size = Pt(11)
    run_n.font.color.rgb = NAVY
    run_n.font.name = 'Calibri'
    run_l = p.add_run(label)
    run_l.font.size = Pt(11)
    run_l.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)
    run_l.font.name = 'Calibri'

# ═══════════════════════════════════════════
#  SECTION 1: INFORMASI UMUM
# ═══════════════════════════════════════════
doc.add_page_break()

doc.add_heading('1. Informasi Umum & Akses Platform', level=1)
add_horizontal_line(doc)

p = doc.add_paragraph()
run = p.add_run(
    'Kios Lumero adalah platform aplikasi pemesanan makanan berbasis web yang dirancang untuk mendukung '
    'operasional restoran secara digital. Sistem ini mencakup fitur Self-Order (pesan mandiri via tablet di outlet) '
    'dan Online Order (pesan dari mana saja via smartphone/desktop). Seluruh proses transaksi, mulai dari pemilihan '
    'menu hingga pembayaran, berlangsung secara real-time dan seamless tanpa perlu memuat ulang halaman.'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

add_styled_table(doc,
    headers=['Parameter', 'Detail'],
    rows=[
        ['Nama Platform', 'Kios Lumero (Self-Order & Online Order)'],
        ['Perusahaan', 'PT. Lokapedia Sukses Bersama'],
        ['Domain Produksi', 'https://lokapedia.id/lumero'],
        ['Halaman Pemesanan', 'https://lokapedia.id/lumero/member/online-order.php'],
        ['Lingkungan Pengujian', 'Midtrans Sandbox Environment'],
        ['Integrasi Midtrans', 'Core API v2 — Payment Type: QRIS'],
        ['Webhook Endpoint', 'https://lokapedia.id/lumero/api/midtrans/notification'],
    ],
    col_widths=[5, 12]
)

# ═══════════════════════════════════════════
#  SECTION 2: KREDENSIAL AKUN UJI COBA
# ═══════════════════════════════════════════
doc.add_page_break()

doc.add_heading('2. Kredensial Akun Uji Coba (Dummy)', level=1)
add_horizontal_line(doc)

p = doc.add_paragraph()
run = p.add_run(
    'Untuk memverifikasi bahwa pesanan yang telah dibayar berhasil masuk ke dalam sistem POS (Point of Sale) '
    'dan ditampilkan di layar Kasir/Dapur, Tim Verifikator dapat melakukan login ke dalam dashboard menggunakan '
    'akun uji coba berikut:'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

# Credential box
cred_table = doc.add_table(rows=5, cols=2)
cred_table.alignment = WD_TABLE_ALIGNMENT.CENTER

cred_data = [
    ('URL Login', 'https://lokapedia.id/lumero/login.php'),
    ('Username', 'midtrans_tester'),
    ('Password', 'Password123!'),
    ('Hak Akses', 'Administrator (Admin Outlet Cabang)'),
    ('Outlet', 'Outlet Midtrans (Verifikasi) — Kode: MIDTRANS'),
]

for i, (label, value) in enumerate(cred_data):
    cell_l = cred_table.rows[i].cells[0]
    cell_v = cred_table.rows[i].cells[1]
    
    cell_l.text = ''
    p = cell_l.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = NAVY
    set_cell_shading(cell_l, 'F1F5F9')
    cell_l.width = Cm(4)
    
    cell_v.text = ''
    p = cell_v.paragraphs[0]
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if i in [1, 2]:
        run.bold = True
        run.font.color.rgb = RED
    cell_v.width = Cm(13)

p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(12)
run = p_note.add_run('Catatan: ')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
run2 = p_note.add_run(
    'Akun ini merupakan akun non-email yang dibuat khusus untuk keperluan verifikasi pengujian oleh Tim Midtrans. '
    'Akun ini memiliki hak akses setingkat Administrator pada outlet cabang khusus bernama "Outlet Midtrans (Verifikasi)" '
    'yang dibuat secara terpisah dari outlet operasional utama. Dengan akun ini, Tim Verifikator dapat melihat pesanan '
    'yang masuk, memantau status pembayaran, serta memverifikasi bahwa alur transaksi berjalan dengan baik.'
)
run2.font.size = Pt(10)
run2.font.color.rgb = GRAY
run2.font.name = 'Calibri'

# ═══════════════════════════════════════════
#  SECTION 3: DIAGRAM USER FLOW
# ═══════════════════════════════════════════
doc.add_page_break()

doc.add_heading('3. Diagram Alur Transaksi (User Flow)', level=1)
add_horizontal_line(doc)

p = doc.add_paragraph()
run = p.add_run(
    'Diagram berikut menggambarkan keseluruhan alur transaksi pembelian online di platform Kios Lumero, '
    'mulai dari akses halaman awal hingga pesanan masuk ke dapur/POS outlet:'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

# Insert flow diagram image
if os.path.exists(FLOW_DIAGRAM):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(FLOW_DIAGRAM, width=Inches(5.5))

p_caption = doc.add_paragraph()
p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_caption.add_run('Gambar 1: Diagram Alur Transaksi Pembelian Online — Kios Lumero')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

# ═══════════════════════════════════════════
#  SECTION 4: DETAIL LANGKAH
# ═══════════════════════════════════════════
doc.add_page_break()

doc.add_heading('4. Detail Langkah-Langkah Pemesanan', level=1)
add_horizontal_line(doc)

p = doc.add_paragraph()
run = p.add_run(
    'Berikut adalah uraian detail dari setiap langkah pemesanan (dari sudut pandang pelanggan), '
    'dilengkapi dengan screenshot untuk mempermudah pemahaman Tim Verifikator:'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

# ──── STEP 1 ────
add_step_heading(doc, 1, 'Akses Halaman Welcome')
p = doc.add_paragraph()
run = p.add_run(
    'Pelanggan mengakses tautan utama Kios Lumero melalui browser pada perangkat smartphone atau desktop. '
    'Sistem akan menampilkan halaman Welcome Screen yang berisi video promosi singkat mengenai brand Lumero '
    'serta tombol interaktif bertuliskan "Mulai Online Order" untuk memulai proses pemesanan.'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

add_screenshot_placeholder(doc, 1,
    'Instruksi: Buka URL https://lokapedia.id/lumero/member/welcome.php di browser.\n'
    'Screenshot seluruh halaman yang menampilkan video promosi dan tombol "Mulai Online Order".'
)

# ──── STEP 2 ────
add_step_heading(doc, 2, 'Deteksi & Pemilihan Cabang (Outlet)')
p = doc.add_paragraph()
run = p.add_run(
    'Saat pelanggan menekan tombol "Mulai Online Order", sistem secara otomatis meminta izin akses GPS '
    '(geolocation) dari perangkat pengguna. Menggunakan formula Haversine, sistem menghitung jarak pelanggan '
    'ke seluruh cabang Kios Lumero yang terdaftar dan aktif, kemudian merekomendasikan cabang terdekat di '
    'urutan paling atas. Pelanggan tinggal memilih cabang yang dituju untuk melanjutkan.'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

add_screenshot_placeholder(doc, 2,
    'Instruksi: Klik tombol "Mulai Online Order" pada halaman Welcome.\n'
    'Screenshot halaman pemilihan cabang yang muncul (select-branch.php), pastikan terlihat daftar cabang\n'
    'beserta indikator jarak dan tombol "Pesan di Sini".'
)

# ──── STEP 3 ────
add_step_heading(doc, 3, 'Pemilihan Menu & Keranjang (Cart)')
p = doc.add_paragraph()
run = p.add_run(
    'Setelah memilih cabang, pelanggan diarahkan ke halaman katalog produk utama. Halaman ini menampilkan '
    'seluruh kategori menu beserta produk yang tersedia (lengkap dengan gambar, harga, dan status ketersediaan). '
    'Pelanggan dapat memilih produk, menentukan varian (misalnya: ukuran, level pedas, dengan/tanpa nasi), '
    'dan menambahkannya ke keranjang belanja. Jumlah item di keranjang akan selalu ter-update secara real-time.'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

add_screenshot_placeholder(doc, 3,
    'Instruksi: Setelah memilih cabang, Anda akan masuk ke halaman online-order.php.\n'
    'Tambahkan 1-2 produk ke keranjang, lalu screenshot layar yang menunjukkan katalog menu\n'
    'dan indikator keranjang (jumlah item & total harga) di bagian bawah layar.'
)

# ──── STEP 4 ────
add_step_heading(doc, 4, 'Checkout & Konfigurasi Logistik')
p = doc.add_paragraph()
run = p.add_run(
    'Pelanggan menekan tombol "Lanjut Bayar" untuk membuka drawer/panel checkout. Di panel ini, pelanggan '
    'melengkapi data diri berupa Nama dan Nomor WhatsApp, kemudian memilih metode pengambilan pesanan:\n\n'
    '• Ambil di Outlet (Pickup): Pesanan disiapkan dan diambil langsung oleh pelanggan di cabang.\n'
    '• Delivery (Diantar Kurir): Pelanggan menentukan titik pengantaran di peta interaktif (Maps). '
    'Ongkos kirim dihitung secara otomatis dan dinamis berdasarkan jarak titik antar ke cabang yang dipilih '
    'sebelumnya menggunakan model perhitungan per-kilometer.'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

add_screenshot_placeholder(doc, 4,
    'Instruksi: Klik tombol "Lanjut Bayar" di bagian bawah halaman online-order.php.\n'
    'Screenshot panel/drawer checkout yang muncul dari bawah layar.\n'
    'Pastikan terlihat kolom Nama, WhatsApp, opsi Pickup/Delivery, dan pilihan metode pembayaran.'
)

# ──── STEP 5 ────
add_step_heading(doc, 5, 'Pembayaran QRIS via Midtrans (Sandbox)')
p = doc.add_paragraph()
run = p.add_run(
    'Pelanggan memilih metode pembayaran "QRIS / E-Wallet" dan menekan tombol "Bayar Sekarang". '
    'Sistem backend Kios Lumero akan berkomunikasi secara real-time dengan Midtrans Core API v2 '
    '(endpoint: /v2/charge) untuk men-generate kode QR QRIS dinamis. '
    'Kode QR yang berhasil di-generate kemudian ditampilkan di layar pelanggan dalam bentuk pop-up overlay '
    'yang elegan, lengkap dengan informasi nominal pembayaran, nomor pesanan, dan hitungan mundur masa berlaku QRIS.'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

add_screenshot_placeholder(doc, 5,
    'Instruksi: Setelah mengisi data di panel checkout, pastikan metode pembayaran "QRIS / E-Wallet" terpilih,\n'
    'lalu klik tombol "Bayar Sekarang".\n'
    'Screenshot pop-up overlay yang muncul di tengah layar, yang berisi gambar QR Code Midtrans,\n'
    'nominal pembayaran, nomor pesanan, dan countdown timer.'
)

# ──── STEP 6 ────
add_step_heading(doc, 6, 'Simulasi Pembayaran (Sandbox)')
p = doc.add_paragraph()
run = p.add_run(
    'Dalam lingkungan Sandbox, Tim Verifikator Midtrans dapat melakukan simulasi pembayaran QRIS dengan cara:\n\n'
    '1. Klik kanan pada gambar QR Code yang ditampilkan di pop-up Kios Lumero.\n'
    '2. Pilih "Copy Image Address" (Salin Alamat Gambar).\n'
    '3. Buka Midtrans QRIS Simulator di tab baru: https://simulator.sandbox.midtrans.com/qris/index\n'
    '4. Tempelkan (paste) URL gambar QR Code ke dalam kolom yang tersedia.\n'
    '5. Pilih Issuer (misalnya: GoPay) dan klik tombol "Pay".\n\n'
    'Dalam skenario Production, pelanggan cukup memindai QR Code ini menggunakan aplikasi E-Wallet '
    'pilihan mereka (GoPay, OVO, Dana, ShopeePay, atau aplikasi M-Banking yang mendukung QRIS).'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

add_screenshot_placeholder(doc, 6,
    'Instruksi: Buka tab baru ke https://simulator.sandbox.midtrans.com/qris/index\n'
    'Paste URL gambar QR Code dari pop-up Kios Lumero, pilih Issuer "GoPay", lalu klik "Pay".\n'
    'Screenshot halaman simulator yang menunjukkan status transaksi berhasil (successful).'
)

# ──── STEP 7 ────
add_step_heading(doc, 7, 'Notifikasi Sukses Secara Real-time')
p = doc.add_paragraph()
run = p.add_run(
    'Begitu pembayaran berhasil diproses (baik melalui simulator maupun scan langsung), sistem Kios Lumero '
    'secara otomatis dan real-time mendeteksi perubahan status transaksi melalui mekanisme polling ke endpoint '
    'check-qris-status.php. Pop-up QR Code akan langsung tertutup dan digantikan oleh pop-up konfirmasi berwarna '
    'hijau bertuliskan "Pembayaran Berhasil!" yang dilengkapi progress bar animasi. '
    'Setelah 3 detik, pelanggan secara otomatis dialihkan ke halaman Lacak Pesanan (Struk Digital) yang menampilkan '
    'detail lengkap pesanan beserta status antrian.'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

add_screenshot_placeholder(doc, 7,
    'Instruksi: Kembali ke tab Kios Lumero segera setelah pembayaran berhasil di simulator.\n'
    'Screenshot pop-up hijau bertuliskan "Pembayaran Berhasil!" beserta progress bar dan nomor pesanan.\n'
    '(Catatan: Pop-up ini hanya tampil selama ±3 detik sebelum redirect, siapkan screenshot cepat).'
)

# ═══════════════════════════════════════════
#  SECTION 5: BACKEND & WEBHOOK
# ═══════════════════════════════════════════
doc.add_page_break()

doc.add_heading('5. Alur Verifikasi Backend & Webhook', level=1)
add_horizontal_line(doc)

p = doc.add_paragraph()
run = p.add_run(
    'Untuk memastikan keandalan dan keamanan sistem, Kios Lumero telah mengimplementasikan mekanisme '
    'penerimaan notifikasi (Webhook / HTTP Notification) sesuai dengan standar integrasi Midtrans. '
    'Berikut adalah alur teknis yang terjadi di sisi backend saat transaksi berhasil:'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

webhook_steps = [
    ('Penerimaan Notifikasi',
     'Saat status transaksi berubah (misalnya dari "pending" menjadi "settlement"), '
     'server Midtrans akan mengirimkan HTTP POST request ke endpoint webhook Kios Lumero di:\n'
     'https://lokapedia.id/lumero/api/midtrans/notification'),
    ('Validasi Signature Key',
     'Sistem Kios Lumero memverifikasi keaslian data notifikasi dengan cara menghitung ulang '
     'signature_key menggunakan format standar Midtrans (SHA-512) dari kombinasi: order_id + status_code '
     '+ gross_amount + server_key. Jika signature tidak cocok, notifikasi akan ditolak secara otomatis.'),
    ('Pembaruan Status Pesanan',
     'Setelah validasi berhasil, status pesanan di database Kios Lumero akan diperbarui secara otomatis '
     'dari "unpaid" menjadi "paid", beserta pencatatan transaction_id dan settlement_time dari Midtrans.'),
    ('Tampilan Dapur / POS',
     'Pesanan yang telah berstatus "paid" secara otomatis muncul di layar aplikasi Kasir/POS '
     'cabang terkait (berdasarkan outlet_id yang dipilih pelanggan saat checkout). '
     'Fitur auto-refresh pada halaman POS memastikan pesanan langsung terlihat tanpa perlu '
     'memuat ulang halaman secara manual.'),
]

for i, (title, desc) in enumerate(webhook_steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run_n = p.add_run(f'{i}. {title}\n')
    run_n.bold = True
    run_n.font.size = Pt(11)
    run_n.font.color.rgb = NAVY
    run_n.font.name = 'Calibri'
    run_d = p.add_run(desc)
    run_d.font.size = Pt(11)
    run_d.font.name = 'Calibri'

add_horizontal_line(doc)

# Verification screenshot
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Bukti Verifikasi Pesanan di Sistem POS')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

p = doc.add_paragraph()
run = p.add_run(
    'Untuk membuktikan bahwa pesanan yang dibayar melalui QRIS Midtrans benar-benar masuk ke sistem operasional '
    'restoran, Tim Verifikator dapat login menggunakan akun uji coba yang tercantum di Bab 2, kemudian '
    'mengakses halaman POS/Kasir untuk melihat pesanan yang baru masuk.'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

add_screenshot_placeholder(doc, 8,
    'Instruksi: Buka tab baru, akses https://lokapedia.id/lumero/login.php\n'
    'Login menggunakan akun midtrans_tester / Password123!\n'
    'Masuk ke halaman POS/Kasir, lalu screenshot daftar pesanan yang menunjukkan\n'
    'orderan terbaru dari uji coba tadi berstatus "Lunas" / "Sudah Dibayar".'
)

# ═══════════════════════════════════════════
#  SECTION 6: PENUTUP
# ═══════════════════════════════════════════
doc.add_page_break()

doc.add_heading('6. Penutup', level=1)
add_horizontal_line(doc)

p = doc.add_paragraph()
run = p.add_run(
    'Demikian dokumen penjelasan alur transaksi pembelian online dan panduan uji coba pembayaran QRIS '
    'pada platform Kios Lumero. Integrasi pembayaran QRIS Midtrans pada sistem kami telah dirancang untuk '
    'berjalan secara real-time, seamless, dan sesuai dengan standar keamanan yang ditetapkan oleh Midtrans.\n\n'
    'Kami sangat mengapresiasi waktu dan perhatian Tim Verifikator Midtrans dalam melakukan peninjauan terhadap '
    'sistem kami. Besar harapan kami agar fasilitas transaksi Production dapat segera diaktifkan sehingga '
    'pelanggan Kios Lumero dapat menikmati pengalaman pembayaran digital yang aman dan nyaman.\n\n'
    'Apabila terdapat pertanyaan, kendala teknis, atau informasi tambahan yang diperlukan selama proses '
    'verifikasi, silakan menghubungi kami melalui:'
)
run.font.size = Pt(11)
run.font.name = 'Calibri'

contact_table = add_styled_table(doc,
    headers=['Kontak', 'Detail'],
    rows=[
        ['Nama PIC', 'Moch. Adriq Fadillah'],
        ['Email', 'mochfadillah1208@gmail.com'],
        ['WhatsApp', '(silakan isi nomor WA aktif)'],
        ['Posisi', 'Developer — PT. Lokapedia Sukses Bersama'],
    ],
    col_widths=[5, 12]
)

p_closing = doc.add_paragraph()
p_closing.paragraph_format.space_before = Pt(24)
p_closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_closing.add_run('Terima kasih atas kerja sama yang baik.\n\n')
run.font.size = Pt(11)
run.font.name = 'Calibri'
run2 = p_closing.add_run('Hormat kami,\nTim Kios Lumero')
run2.bold = True
run2.font.size = Pt(11)
run2.font.color.rgb = NAVY
run2.font.name = 'Calibri'

# ──── SAVE ────
doc.save(OUTPUT_FILE)
print(f'[OK] Dokumen berhasil di-generate: {os.path.abspath(OUTPUT_FILE)}')
